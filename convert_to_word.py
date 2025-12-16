#!/usr/bin/env python3
"""
Convert markdown analysis to Word document
"""

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("Installing required package...")
    import subprocess
    subprocess.check_call(['pip3', 'install', 'python-docx'])
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

import re

# Read the markdown file
with open('/Users/macbookairm3/.gemini/antigravity/brain/319fc594-671a-4f1e-ab68-2f68889e60bd/implemented_features_analysis.md', 'r') as f:
    content = f.read()

# Create a new Document
doc = Document()

# Set document margins
sections = doc.sections
for section in sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

def add_heading(text, level=1):
    """Add a heading to the document"""
    heading = doc.add_heading(text, level=level)
    return heading

def add_paragraph(text, style=None):
    """Add a paragraph to the document"""
    if style:
        p = doc.add_paragraph(text, style=style)
    else:
        p = doc.add_paragraph(text)
    return p

def parse_markdown_table(table_text):
    """Parse a markdown table and add it to the document"""
    lines = [line.strip() for line in table_text.strip().split('\n') if line.strip()]
    if len(lines) < 2:
        return
    
    # Extract headers
    headers = [cell.strip() for cell in lines[0].split('|') if cell.strip()]
    
    # Skip separator line
    # Extract data rows
    rows = []
    for line in lines[2:]:
        cells = [cell.strip() for cell in line.split('|') if cell.strip()]
        if cells:
            rows.append(cells)
    
    if not rows:
        return
    
    # Create table
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    
    # Add headers
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        if i < len(header_cells):
            header_cells[i].text = header
            # Make header bold
            for paragraph in header_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.bold = True
    
    # Add data rows
    for row_idx, row_data in enumerate(rows):
        row_cells = table.rows[row_idx + 1].cells
        for col_idx, cell_data in enumerate(row_data):
            if col_idx < len(row_cells):
                row_cells[col_idx].text = cell_data

# Process the markdown content
lines = content.split('\n')
i = 0
in_table = False
table_buffer = []

while i < len(lines):
    line = lines[i]
    
    # Check for table start
    if '|' in line and not in_table:
        in_table = True
        table_buffer = [line]
        i += 1
        continue
    
    # Collect table lines
    if in_table:
        if '|' in line:
            table_buffer.append(line)
            i += 1
            continue
        else:
            # End of table
            in_table = False
            parse_markdown_table('\n'.join(table_buffer))
            table_buffer = []
            continue
    
    # Heading 1
    if line.startswith('# '):
        add_heading(line[2:], level=1)
    # Heading 2
    elif line.startswith('## '):
        add_heading(line[3:], level=2)
    # Heading 3
    elif line.startswith('### '):
        add_heading(line[4:], level=3)
    # Heading 4
    elif line.startswith('#### '):
        add_heading(line[5:], level=4)
    # Horizontal rule
    elif line.strip() == '---':
        doc.add_paragraph('_' * 80)
    # Bullet point
    elif line.strip().startswith('- ') or line.strip().startswith('* '):
        text = line.strip()[2:]
        # Remove markdown formatting
        text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)  # Bold
        text = re.sub(r'\*(.*?)\*', r'\1', text)  # Italic
        text = re.sub(r'`(.*?)`', r'\1', text)  # Code
        add_paragraph(text, style='List Bullet')
    # Numbered list
    elif re.match(r'^\d+\.\s', line.strip()):
        text = re.sub(r'^\d+\.\s', '', line.strip())
        # Remove markdown formatting
        text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
        text = re.sub(r'\*(.*?)\*', r'\1', text)
        text = re.sub(r'`(.*?)`', r'\1', text)
        add_paragraph(text, style='List Number')
    # Regular paragraph
    elif line.strip():
        text = line.strip()
        # Remove markdown formatting
        text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
        text = re.sub(r'\*(.*?)\*', r'\1', text)
        text = re.sub(r'`(.*?)`', r'\1', text)
        # Remove links but keep text
        text = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', text)
        
        # Check for special markers
        if text.startswith('✅'):
            p = add_paragraph(text)
            p.runs[0].font.color.rgb = RGBColor(0, 128, 0)
        elif text.startswith('❌'):
            p = add_paragraph(text)
            p.runs[0].font.color.rgb = RGBColor(255, 0, 0)
        elif text.startswith('⚠️'):
            p = add_paragraph(text)
            p.runs[0].font.color.rgb = RGBColor(255, 165, 0)
        elif text.startswith('🆕'):
            p = add_paragraph(text)
            p.runs[0].font.color.rgb = RGBColor(0, 0, 255)
        else:
            add_paragraph(text)
    
    i += 1

# Save the document
output_path = '/Users/macbookairm3/Desktop/Uni/Final Project Stack/institute-management-system/Implemented_Features_Analysis.docx'
doc.save(output_path)
print(f"✅ Word document created successfully: {output_path}")
